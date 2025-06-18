
import pandas as pd
import numpy as np
from docx import Document

def analyse_enquête_met_weging(df, kolomnamen, kolom_labels, gewichtsvariabele="weging"):
    analyseresultaten = []
    gewogen = gewichtsvariabele in df.columns

    for kolom in kolomnamen:
        if kolom not in df.columns:
            continue

        vraagtekst = kolom_labels.get(kolom, kolom)

        if pd.api.types.is_numeric_dtype(df[kolom]) and df[kolom].nunique() > 10:
            if gewogen:
                gew = df[kolom] * df[gewichtsvariabele]
                gemiddelde = gew.sum() / df[gewichtsvariabele].sum()
            else:
                gemiddelde = df[kolom].mean()

            analyseresultaten.append({
                "vraag": f"{kolom}: {vraagtekst}",
                "type": "schaal",
                "data": {
                    "gemiddelde": gemiddelde
                }
            })
        else:
            if gewogen:
                gewogen_freq = df.groupby(kolom)[gewichtsvariabele].sum()
                totaal = gewogen_freq.sum()
                percentages = (gewogen_freq / totaal * 100).round(1)
            else:
                percentages = (df[kolom].value_counts(normalize=True) * 100).round(1)

            analyseresultaten.append({
                "vraag": f"{kolom}: {vraagtekst}",
                "type": "frequentie",
                "data": list(percentages.items())
            })

    return analyseresultaten, gewogen

def schrijf_resultaten_naar_word(analyseresultaten, bestandsnaam="rapport_enquete.docx", gewogen=True):
    doc = Document()
    doc.add_heading("Samenvatting enquête-analyse", level=1)

    if gewogen:
        doc.add_paragraph("⚖️ De onderstaande resultaten zijn gewogen op basis van de variabele 'weging'.")
    else:
        doc.add_paragraph("⚠️ De onderstaande resultaten zijn ongewogen.")

    for resultaat in analyseresultaten:
        vraag = resultaat["vraag"]
        doc.add_heading(vraag, level=2)

        if resultaat["type"] == "frequentie":
            table = doc.add_table(rows=1, cols=2)
            table.autofit = True
            hdr_cells = table.rows[0].cells
            hdr_cells[0].text = "Antwoordoptie"
            hdr_cells[1].text = "Percentage"

            for label, percentage in resultaat["data"]:
                row_cells = table.add_row().cells
                row_cells[0].text = str(label)
                row_cells[1].text = f"{percentage} %"

        elif resultaat["type"] == "schaal":
            schaal = resultaat["data"]
            doc.add_paragraph(f"Gemiddelde: {schaal['gemiddelde']:.1f}")

        doc.add_paragraph("")  # lege regel tussen vragen

    output_path = f"/mnt/data/enquete_analyse_functies_exec_ready.py"
    with open(output_path, "w") as file:
        file.write('')
    doc.save(output_path)
    return output_path
